"""Generate LCHAI Questionnaire as a well-formatted Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT_DIR = r"D:\Dropbox\PHD\THESIS\SLIMA_Thesis_Research_ver_2026_rev_9\appendix"

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

for s in doc.styles:
    if s.name and s.name.startswith('Heading'):
        s.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_likert_block(questions):
    """Add a Likert-scale question block. Each question gets its own row
    with the full question text spanning the width, and a second row with
    5 Likert checkboxes."""
    table = doc.add_table(rows=0, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr = table.add_row().cells
    hdr[0].text = "#"
    hdr[1].text = "Pregunta"
    for i, lbl in enumerate(["1", "2", "3", "4", "5"]):
        hdr[2 + i].text = lbl
    for cell in hdr:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
        shading = cell._element.get_or_add_tcPr()
        bg = qn('w:shd')
        s = shading.makeelement(bg, {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'DBEAFE',
        })
        shading.append(s)

    for qid, qtext in questions:
        row = table.add_row().cells
        row[0].text = qid
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for p in row[0].paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)

        row[1].text = qtext
        for p in row[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

        for i in range(5):
            row[2 + i].text = "☐"
            row[2 + i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in row[2 + i].paragraphs[0].runs:
                r.font.size = Pt(10)

    col_widths = [Cm(1.2), Cm(10.5), Cm(1.0), Cm(1.0), Cm(1.0), Cm(1.0), Cm(1.0)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════

add_heading("LCHAI v2.0 — Cuestionario de Evaluación de Explicabilidad", 1)
add_para("Lung Cancer Histologic Analysis with AI", bold=True, size=12)
add_para("Evaluación por expertos de SOLCA Ecuador", italic=True)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Investigador: ").bold = True
p.add_run("Servio F. Lima Reina")
p = doc.add_paragraph()
p.add_run("Universidad: ").bold = True
p.add_run("University of Fribourg, Switzerland")
p = doc.add_paragraph()
p.add_run("Sistema: ").bold = True
p.add_run("LCHAI v2.0 — https://lchai.gptfy.biz")
p = doc.add_paragraph()
p.add_run("Fecha: ").bold = True
p.add_run("________________")

# ═══════ INSTRUCCIONES ═══════

add_heading("Instrucciones", 2)
add_para("Usted ha sido invitado(a) a evaluar la herramienta LCHAI v2.0, un sistema de investigación basado en inteligencia artificial para el análisis histológico de adenocarcinoma de pulmón (LUAD).")
doc.add_paragraph()
add_para("Por favor:", bold=True)
doc.add_paragraph("Revise al menos 2 casos TCGA pre-cargados en el sistema", style='List Number')
doc.add_paragraph("Navegue por las pestañas: Analysis, Viewer, Graph, y Admin", style='List Number')
doc.add_paragraph("En la pestaña Analysis, explore al menos 2 sub-tabs de genes diferentes", style='List Number')
doc.add_paragraph("Responda cada pregunta usando la escala de 1 a 5", style='List Number')

doc.add_paragraph()
add_para("IMPORTANTE: LCHAI es una herramienta de investigación, NO un dispositivo de diagnóstico clínico.", bold=True)

doc.add_paragraph()
add_para("Escala de valoración:", bold=True)
scale_table = doc.add_table(rows=6, cols=2)
scale_table.style = 'Table Grid'
scale_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = [("Puntuación", "Significado"),
           ("1", "Totalmente en desacuerdo / Nada útil"),
           ("2", "En desacuerdo / Poco útil"),
           ("3", "Neutral / Moderadamente útil"),
           ("4", "De acuerdo / Útil"),
           ("5", "Totalmente de acuerdo / Muy útil")]
for i, (a, b) in enumerate(headers):
    scale_table.rows[i].cells[0].text = a
    scale_table.rows[i].cells[1].text = b
    scale_table.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in scale_table.rows[i].cells[0].paragraphs[0].runs:
        r.bold = True

# ═══════ SECCIÓN E ═══════

add_heading("Sección E: Datos del Evaluador", 2)
for label in [
    "E1. Nombre (opcional): ___________________________",
    "E2. Rol profesional:  ☐ Histopatólogo  ☐ Biólogo molecular  ☐ Oncólogo  ☐ Otro: _______",
    "E3. Años de experiencia:  ☐ <5  ☐ 5-10  ☐ 10-20  ☐ >20",
    "E4. Experiencia con IA en patología:  ☐ Ninguna  ☐ Básica  ☐ Intermedia  ☐ Avanzada",
    "E5. Institución:  ☐ SOLCA Guayaquil  ☐ SOLCA Quito  ☐ HFR Suiza  ☐ Otra: _______",
]:
    doc.add_paragraph(label)

# ═══════ SECCIÓN A ═══════

add_heading("Sección A: Evaluación General (todos los perfiles)", 2)
add_para("Revise la pestaña Analysis y seleccione al menos 2 genes en los sub-tabs.", italic=True, size=9)

add_heading("A1. Usabilidad", 3)
add_likert_block([
    ("A1.1", "La interfaz de LCHAI es fácil de navegar y entender"),
    ("A1.2", "La organización por sub-tabs de genes (ordenados por P(mut)) es intuitiva"),
    ("A1.3", "El flujo de trabajo (subir imagen → analizar → explorar por gen) es claro"),
    ("A1.4", "La información presentada por gen es suficiente para evaluar el caso"),
])

add_heading("A2. Confianza", 3)
add_likert_block([
    ("A2.1", "El sistema me genera confianza como herramienta de apoyo a la decisión"),
    ("A2.2", "Las etiquetas 'Conclusive' / 'Inconclusive' me ayudan a saber cuánto confiar en cada predicción"),
    ("A2.3", "El disclaimer 'Research tool — NOT for clinical diagnosis' es visible y apropiado"),
    ("A2.4", "Las referencias a PubMed [1], [2] en la explicación incrementan mi confianza en el sistema"),
])

add_heading("A3. Explicación Generada por IA", 3)
add_likert_block([
    ("A3.1", "La explicación auto-generada por gen es comprensible"),
    ("A3.2", "La explicación es médicamente correcta (según mi conocimiento)"),
    ("A3.3", "Las citas bibliográficas [1], [2] con links a PubMed son útiles y verificables"),
    ("A3.4", "La explicación utiliza términos de intensidad apropiados (ej. 'Balanced', 'Moderate Synergy')"),
])

# ═══════ SECCIÓN B ═══════

add_heading("Sección B: Histopatólogos", 2)
add_para("Revise los overlays de patrones y atención. Pase el mouse sobre las imágenes para ver los tooltips.", italic=True, size=9)

add_heading("B1. Overlay de Patrones", 3)
add_likert_block([
    ("B1.1", "Los colores del overlay de patrones son distinguibles entre sí"),
    ("B1.2", "Las regiones coloreadas corresponden a tejido real (no fondo ni artefactos)"),
    ("B1.3", "El patrón asignado a cada región coincide con mi evaluación visual del H&E"),
    ("B1.4", "La clasificación del patrón predominante es correcta para este slide"),
    ("B1.5", "Los porcentajes de composición en la leyenda son razonables"),
    ("B1.6", "El tooltip interactivo al pasar el mouse (nombre del patrón y %) es útil"),
])

add_heading("B2. Mapa de Atención (ABMIL)", 3)
add_likert_block([
    ("B2.1", "Las regiones destacadas por el mapa de atención son clínicamente relevantes"),
    ("B2.2", "El mapa de atención me ayuda a entender DÓNDE el modelo concentra su análisis"),
    ("B2.3", "Las regiones de alta atención coinciden con áreas diagnósticamente importantes"),
    ("B2.4", "El tooltip fuzzy de atención (ej. 'High Attention, p97') es informativo"),
])

add_heading("B3. Vista lado a lado", 3)
add_likert_block([
    ("B3.1", "Ver overlay de patrones y atención simultáneamente es más útil que por separado"),
    ("B3.2", "Puedo correlacionar visualmente qué patrones reciben mayor atención del modelo"),
])

add_heading("B4. Corrección de Patrones (Active Learning)", 3)
add_likert_block([
    ("B4.1", "La herramienta de lasso para corregir patrones es fácil de usar"),
    ("B4.2", "Es útil poder corregir la clasificación de patrones del modelo"),
    ("B4.3", "Estaría dispuesto(a) a usar esta herramienta regularmente"),
])

# ═══════ SECCIÓN C ═══════

add_heading("Sección C: Biólogos Moleculares", 2)
add_para("Revise la descomposición SHAP, la comparación de ablación, y la sección Choquet (si aparece).", italic=True, size=9)

add_heading("C1. Descomposición SHAP con Etiqueta Fuzzy", 3)
add_likert_block([
    ("C1.1", "El gráfico SHAP (embedding vs. pattern) es fácil de interpretar"),
    ("C1.2", "La etiqueta fuzzy (ej. 'Embedding-Led', 'Balanced') me ayuda a interpretar el balance"),
    ("C1.3", "La descomposición SHAP me ayuda a entender si el modelo usa morfología visible o features sub-celulares"),
])

add_heading("C2. Comparación de Ablación (3 modelos)", 3)
add_likert_block([
    ("C2.1", "Las 3 barras (Combined, Emb-only, Pat-only) me permiten entender la contribución de cada feature"),
    ("C2.2", "La nota aclaratoria ('Each bar shows P(mut) from a different model...') evita malinterpretaciones"),
    ("C2.3", "El indicador delta (+X pp o -X pp) me dice claramente si los patrones ayudan o interfieren"),
])

add_heading("C3. Choquet Shapley Values (condicional)", 3)
add_para("Seleccione un gen donde la sección Choquet sea visible (ej. KRAS).", italic=True, size=9)
add_likert_block([
    ("C3.1", "El badge de perfil Shapley (ej. 'Uniform', 'Near-Uniform') da una idea rápida de preferencia"),
    ("C3.2", "Las etiquetas fuzzy por patrón (ej. 'Average', 'Slightly Above') son más informativas que solo el número"),
    ("C3.3", "Las clasificaciones de interacciones (ej. 'Moderate Synergy ↑') ayudan a interpretar sin ser experto"),
    ("C3.4", "Las sinergias identificadas (ej. solid × lepidic para KRAS) son biológicamente plausibles"),
    ("C3.5", "Que la sección Choquet solo aparezca cuando los patrones contribuyen positivamente es apropiado"),
])

add_heading("C4. Etiquetas Fuzzy Lingüísticas (general)", 3)
add_likert_block([
    ("C4.1", "Las etiquetas lingüísticas hacen los resultados numéricos más accesibles"),
    ("C4.2", "Confío en que las etiquetas de la explicación IA corresponden a lo que veo en los gráficos"),
    ("C4.3", "La posibilidad de ajustar los umbrales (Admin → Fuzzy Labels) es valiosa"),
])

# ═══════ SECCIÓN D ═══════

add_heading("Sección D: Oncólogos", 2)
add_para("Revise la pestaña Graph (Knowledge Graph) y el flujo clínico completo.", italic=True, size=9)

add_heading("D1. Predicción de Mutación", 3)
add_likert_block([
    ("D1.1", "Las probabilidades de mutación son útiles para priorizar tests moleculares"),
    ("D1.2", "El método de predicción (B2/P/FC) está claramente indicado"),
    ("D1.3", "La información de AUROC y confianza me permite evaluar la fiabilidad"),
])

add_heading("D2. Grafo de Conocimiento", 3)
add_likert_block([
    ("D2.1", "El grafo presenta correctamente las asociaciones gen-tratamiento"),
    ("D2.2", "La proveniencia (PMIDs) al hover sobre aristas es útil y verificable"),
    ("D2.3", "El grafo me ayuda a conectar mutación con opciones terapéuticas"),
])

add_heading("D3. Flujo Clínico Integrado", 3)
add_likert_block([
    ("D3.1", "La combinación overlay + SHAP + ablación + explicación proporciona un cuadro completo por gen"),
    ("D3.2", "LCHAI podría reducir el tiempo de espera para decisión de terapia dirigida"),
    ("D3.3", "Recomendaría LCHAI como herramienta complementaria al test molecular"),
    ("D3.4", "Sin test molecular disponible (7-21 días), LCHAI aportaría valor como primera aproximación"),
])

add_heading("D4. Comparación de Explicaciones", 3)
add_likert_block([
    ("D4.1", "Las explicaciones de LCHAI son más útiles que solo un número (ej. 'KRAS: 67%')"),
    ("D4.2", "Las explicaciones visuales (overlay + attention) aportan más que las textuales (IA)"),
    ("D4.3", "La combinación (visual + textual + fuzzy labels) es superior a cualquiera por separado"),
])

# ═══════ COMENTARIOS ═══════

add_heading("Comentarios Abiertos (opcional)", 2)

for q in [
    "¿Qué aspecto de LCHAI le resultó más útil?",
    "¿Las etiquetas fuzzy le ayudaron a interpretar los resultados? ¿Por qué?",
    "¿Qué mejoraría en la herramienta?",
    "¿Algún comentario adicional?",
]:
    add_para(q, bold=True, size=10)
    doc.add_paragraph("_" * 80)
    doc.add_paragraph("_" * 80)
    doc.add_paragraph()

# ═══════ FOOTER ═══════

doc.add_paragraph()
add_para("Gracias por su participación.", bold=True, size=11)
add_para("Este cuestionario es parte de la investigación doctoral 'Lung Cancer Histologic Analysis with AI' (LCHAI). Sus respuestas serán reportadas de forma anónima y agregada.", italic=True, size=9)
add_para("Aprobado por: Comité de Investigación SOLCA-CISOL", italic=True, size=9)

# Save
out_docx = os.path.join(OUT_DIR, "LCHAI_Explainability_Questionnaire_SOLCA.docx")
doc.save(out_docx)
print(f"Saved: {out_docx}")
